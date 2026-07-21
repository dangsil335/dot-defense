# -*- coding: utf-8 -*-
"""
전달용 기획서 docx 빌더.

기존 기획서(친구_시스템_정리.docx)를 템플릿으로 열어 본문만 비우고 채운다.
문단을 새로 만들지 않고 템플릿 문단을 deepcopy 해서 텍스트만 갈아끼우므로
폰트·여백·불릿 numPr·들여쓰기가 전부 그대로 상속된다.

사용:
    doc = Doc()
    doc.h1("문서 제목")
    doc.h2("대분류")
    doc.body("줄글 서술 문단.")
    doc.bullet("불릿 항목")
    doc.table([["헤더1","헤더2"], ["값1","값2"]])
    doc.save(OUT)
"""
import copy
import os
import sys
import tempfile

import docx
from docx.oxml.ns import qn

TEMPLATE = r"C:\Users\solid\Desktop\아이템\시뮬레이터\내용 참고용 요약본\친구_시스템_정리.docx"


class Doc:
    def __init__(self, template=TEMPLATE):
        self.d = docx.Document(template)
        self._grab_samples()
        self._clear_body()

    # --- 템플릿에서 유형별 샘플 문단 확보 -------------------------------
    def _grab_samples(self):
        self.sample = {}
        for p in self.d.paragraphs:
            if not p.text.strip():
                continue
            name = p.style.name if p.style is not None else "_body"
            key = {
                "Heading 1": "h1",
                "Heading 2": "h2",
                "Heading 3": "h3",
                "Heading 4": "h4",
                "List Paragraph": "bullet",
                "_body": "body",
            }.get(name)
            if key and key not in self.sample:
                self.sample[key] = copy.deepcopy(p._element)

        # 템플릿 본문에 h4 샘플이 없으면 h3 을 복제해 pStyle 만 Heading 4 로 바꾼다.
        # (스타일 정의 자체는 문서에 존재하므로 서식은 정상 적용된다.)
        if "h4" not in self.sample and "h3" in self.sample:
            el = copy.deepcopy(self.sample["h3"])
            pPr = el.find(qn("w:pPr"))
            if pPr is None:
                pPr = el.makeelement(qn("w:pPr"), {})
                el.insert(0, pPr)
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is None:
                pStyle = pPr.makeelement(qn("w:pStyle"), {})
                pPr.insert(0, pStyle)
            # styles["Heading 4"] 는 한글 Word 문서에서 KeyError 가 나므로 순회로 찾는다.
            sid = None
            for s in self.d.styles:
                try:
                    if s.type == 1 and s.name == "Heading 4":
                        sid = s.style_id
                        break
                except Exception:
                    continue
            if sid:
                pStyle.set(qn("w:val"), sid)
            self.sample["h4"] = el

        for need in ("h1", "h2", "h3", "body", "bullet"):
            if need not in self.sample:
                raise RuntimeError("템플릿에서 '%s' 샘플 문단을 찾지 못했습니다." % need)

        # 표 샘플 (있으면)
        self.sample_tbl = None
        if self.d.tables:
            self.sample_tbl = copy.deepcopy(self.d.tables[0]._element)

    def _clear_body(self):
        body = self.d.element.body
        for child in list(body):
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)

    # --- 문단 추가 -------------------------------------------------------
    def _add(self, key, text):
        el = copy.deepcopy(self.sample[key])
        runs = el.findall(qn("w:r"))
        # 첫 run 만 남기고 텍스트 교체, 나머지 제거
        if runs:
            first = runs[0]
            for extra in runs[1:]:
                el.remove(extra)
            for t in first.findall(qn("w:t")):
                first.remove(t)
            t = first.makeelement(qn("w:t"), {})
            t.text = text
            t.set(qn("xml:space"), "preserve")
            first.append(t)
        else:
            r = el.makeelement(qn("w:r"), {})
            t = r.makeelement(qn("w:t"), {})
            t.text = text
            r.append(t)
            el.append(r)

        sectPr = self.d.element.body.find(qn("w:sectPr"))
        if sectPr is not None:
            sectPr.addprevious(el)
        else:
            self.d.element.body.append(el)
        return el

    def h1(self, text):
        return self._add("h1", text)

    def h2(self, text):
        return self._add("h2", text)

    def h3(self, text):
        return self._add("h3", text)

    def h4(self, text):
        return self._add("h4", text)

    def body(self, text):
        return self._add("body", text)

    def bullet(self, text):
        return self._add("bullet", text)

    def bullets(self, items):
        for it in items:
            self.bullet(it)

    # --- 표 --------------------------------------------------------------
    def table(self, rows):
        """rows[0] 을 헤더로 보고 표를 만든다. 스타일은 Table Grid."""
        if not rows:
            return None
        t = self.d.add_table(rows=len(rows), cols=len(rows[0]))
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                cell.text = str(val)
                if ri == 0:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.bold = True
        # add_table 은 body 끝(sectPr 뒤)에 붙으므로 sectPr 앞으로 이동
        sectPr = self.d.element.body.find(qn("w:sectPr"))
        if sectPr is not None:
            sectPr.addprevious(t._element)
        return t

    # --- 저장 ------------------------------------------------------------
    def save(self, path):
        try:
            self.d.save(path)
            print("saved:", path)
            return path
        except PermissionError:
            # 대상이 Word 에 열려 있으면 temp 로 우회한다. 절대 삭제하지 않는다.
            alt = os.path.join(tempfile.gettempdir(), os.path.basename(path))
            self.d.save(alt)
            print("PermissionError -> temp 저장:", alt, file=sys.stderr)
            print("   대상 파일이 Word 에서 열려 있습니다. 닫은 뒤 위 파일을 복사해 주세요.", file=sys.stderr)
            return alt
